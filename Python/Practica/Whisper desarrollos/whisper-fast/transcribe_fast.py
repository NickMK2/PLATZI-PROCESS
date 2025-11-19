# main_generador_actas.py
# =======================
# Requisitos:
#   pip install openai-whisper python-docx requests
#   (FFmpeg en PATH)
#
# Menú:
#   1) Transcribir (FIEL 1:1) con Whisper (usa modelo sugerido o manual si está fijado)
#   2) Generar ACTA LITERAL (Word, sin cambios) usando .srt si existe
#   3) Ver configuración detectada (GPU/VRAM y sugerencia de modelo)
#   4) Elegir modelo Whisper por número -> transcribe de inmediato -> preguntar si mantenerlo
#   5) ACTA con IA local: IA SOLO ENCABEZADO + transcripción ÍNTEGRA (verbatim)
#   6) Salir

import os
import sys
import re
import json
import math
import datetime
import shutil
import difflib

# ====== IMPORTS ======
try:
    import torch
except Exception as e:
        print("[AVISO] No se pudo importar torch. Se usará CPU. Detalle:", e)
        torch = None

try:
    import whisper
except Exception as e:
    print("[ERROR] Falta openai-whisper. Instala con: pip install openai-whisper")
    sys.exit(1)

try:
    import requests
except Exception as e:
    print("[ERROR] Falta requests. Instala con: pip install requests")
    sys.exit(1)

try:
    from docx import Document
    from docx.shared import Pt
except Exception as e:
    print("[ERROR] Falta python-docx. Instala con: pip install python-docx")
    sys.exit(1)

# ====== RUTAS / DIRS ======
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
TRANSCRIPCION_DIR = os.path.join(BASE_DIR, "transcripciones")
ACTAS_DIR = os.path.join(BASE_DIR, "actas_generadas")
os.makedirs(TRANSCRIPCION_DIR, exist_ok=True)
os.makedirs(ACTAS_DIR, exist_ok=True)

CONFIG_PATH = os.path.join(BASE_DIR, "config_local.json")  # persiste modelo manual y base_url LM

# ====== REGLAS NOMBRES / LIMPIEZA RUTA ======
ZERO_WIDTHS = ["\u200b", "\u200c", "\u200d", "\ufeff"]  # ZWSP, ZWNJ, ZWJ, BOM
UNICODE_SPACES = [
    "\u00a0",  # NBSP
    "\u202f",  # NNBSP
    "\u2007", "\u2009", "\u2002", "\u2003", "\u2004", "\u2005", "\u2006",
    "\u2008", "\u200a", "\u205f"
]

def _strip_zero_width(s: str) -> str:
    for ch in ZERO_WIDTHS:
        s = s.replace(ch, "")
    return s

def _normalize_spaces(s: str) -> str:
    for ch in UNICODE_SPACES:
        s = s.replace(ch, " ")
    s = re.sub(r"\s+", " ", s)
    return s.strip()

def limpiar_ruta(ruta: str) -> str:
    if not ruta:
        return ruta
    ruta = ruta.strip()
    if ruta.startswith("& "):
        ruta = ruta[2:]
    elif ruta.startswith("&"):
        ruta = ruta[1:]
    ruta = ruta.strip("'\"")
    ruta = _strip_zero_width(ruta)
    ruta = _normalize_spaces(ruta)
    ruta = os.path.normpath(ruta)

    # Quitar espacios justo antes de la extensión ("archivo .m4a" -> "archivo.m4a")
    head, tail = os.path.split(ruta)
    tail = re.sub(r"\s+(\.[A-Za-z0-9]{1,5})$", r"\1", tail)
    return os.path.join(head, tail)

def _norm_name_for_match(name: str) -> str:
    name = _strip_zero_width(name)
    name = _normalize_spaces(name)
    name = re.sub(r"\s+(\.[A-Za-z0-9]{1,5})$", r"\1", name)
    return name.lower()

def resolver_ruta_audio(ruta: str) -> str | None:
    """Si la ruta no existe, intenta encontrar el archivo equivalente en su carpeta."""
    if os.path.isfile(ruta):
        return ruta

    head, tail = os.path.split(ruta)
    if not head or not os.path.isdir(head):
        return None

    target = _norm_name_for_match(tail)
    candidatos = os.listdir(head)
    norm_lista = [_norm_name_for_match(n) for n in candidatos]

    # Coincidencia exacta
    for i, nrm in enumerate(norm_lista):
        if nrm == target:
            return os.path.join(head, candidatos[i])

    # Coincidencia aproximada
    cercano = difflib.get_close_matches(target, norm_lista, n=1, cutoff=0.85)
    if cercano:
        idx = norm_lista.index(cercano[0])
        return os.path.join(head, candidatos[idx])

    return None

def ffmpeg_disponible() -> bool:
    return shutil.which("ffmpeg") is not None

# ====== CONFIG WHISPER / SUGERENCIAS ======
WHISPER_MODELOS = [
    ("tiny",      "~1 GB VRAM / CPU OK",       "Rápido, menor exactitud"),
    ("base",      "~1.5 GB VRAM / CPU OK",     "Mejor que tiny"),
    ("small",     "~3 GB VRAM",                "Buena exactitud"),
    ("medium",    "~7 GB VRAM",                "Alta exactitud"),
    ("large-v3",  "≥12 GB VRAM (GPU)",         "Máxima exactitud"),
]

def _leer_config():
    if os.path.isfile(CONFIG_PATH):
        try:
            with open(CONFIG_PATH, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            return {}
    return {}

def _guardar_config(cfg: dict):
    try:
        with open(CONFIG_PATH, "w", encoding="utf-8") as f:
            json.dump(cfg, f, ensure_ascii=False, indent=2)
    except Exception as e:
        print("[AVISO] No se pudo guardar config:", e)

def _modelo_sugerido_por_vram(vram_gb: float | None) -> str:
    if vram_gb is None:
        return "small"  # CPU: equilibrio
    if vram_gb >= 11.5:
        return "large-v3"
    elif vram_gb >= 7:
        return "medium"
    elif vram_gb >= 3:
        return "small"
    else:
        return "base"

def elegir_config_whisper():
    """Devuelve (model_name, device, fp16, recomendado, vram_gb) según GPU/VRAM y override manual si existe."""
    cfg = _leer_config()
    override = cfg.get("manual_whisper_model") or None

    use_cuda = torch is not None and torch.cuda.is_available()
    device = "cuda" if use_cuda else "cpu"
    fp16 = True if device == "cuda" else False
    vram_gb = None

    if use_cuda:
        try:
            props = torch.cuda.get_device_properties(0)
            vram_gb = props.total_memory / (1024**3)
        except Exception:
            vram_gb = None

    sugerido = _modelo_sugerido_por_vram(vram_gb)
    model_name = override or sugerido
    return (model_name, device, fp16, sugerido, vram_gb)

def mostrar_config():
    print("\n=== DETECCIÓN DE ENTORNO ===")
    print(f"ffmpeg en PATH: {ffmpeg_disponible()}")
    try:
        print(f"torch: {torch.__version__ if torch else 'no-import'}")
        print(f"CUDA disponible: {torch.cuda.is_available() if torch else False}")
        print(f"torch.version.cuda: {getattr(torch.version, 'cuda', None) if torch else None}")
        vram_gb = None
        if torch and torch.cuda.is_available():
            props = torch.cuda.get_device_properties(0)
            print(f"GPU: {torch.cuda.get_device_name(0)}")
            vram_gb = props.total_memory / (1024**3)
            print(f"VRAM: {vram_gb:.1f} GB")
    except Exception as e:
        print("[AVISO] No se pudieron leer propiedades de CUDA:", e)

    model_name, device, fp16, recomendado, vram_gb = elegir_config_whisper()
    print("--- CONFIG WHISPER ---")
    print(f"Modelo en uso: {model_name} ({'manual' if _leer_config().get('manual_whisper_model') else 'sugerido'})")
    print(f"Sugerido para tu hardware: {recomendado}")
    print(f"Device: {device} | fp16={fp16}\n")

# ====== UTIL TRANSCRIPCIÓN ======
def _fmt_ts(sec: float) -> str:
    # HH:MM:SS,mmm
    if sec is None:
        sec = 0.0
    ms = int(round((sec - int(sec)) * 1000))
    s = int(math.floor(sec))
    h = s // 3600
    m = (s % 3600) // 60
    s = s % 60
    return f"{h:02d}:{m:02d}:{s:02d},{ms:03d}"

def _write_srt(segments, out_path):
    with open(out_path, "w", encoding="utf-8") as f:
        for i, seg in enumerate(segments, start=1):
            start = _fmt_ts(seg.get("start"))
            end = _fmt_ts(seg.get("end"))
            text = (seg.get("text") or "").strip()
            f.write(f"{i}\n{start} --> {end}\n{text}\n\n")

def _safe_antibucle(text: str) -> str:
    """DESACTIVADO: NO tocar la transcripción."""
    # return re.sub(r"\b(\w+)(\s+\1){2,}\b", r"\1 \1", text, flags=re.IGNORECASE)
    return text

def _transcribir_fiel(ruta_audio: str, model_name: str, device: str, fp16: bool):
    print("\n--- TRANSCRIPCIÓN FIEL (SIN CAMBIOS) ---")
    print(f"Archivo : {ruta_audio}")
    print(f"Modelo  : {model_name} | device={device} | fp16={fp16}")
    print("Parámetros: temperature=0.0, condition_on_previous_text=False, language='es'")

    model = whisper.load_model(model_name, device=device)

    kwargs = dict(
        language="es",
        task="transcribe",
        temperature=0.0,
        condition_on_previous_text=False,
        no_speech_threshold=0.6,
        logprob_threshold=-1.2,
        compression_ratio_threshold=2.6,
        fp16=fp16,
        verbose=False
    )

    try:
        result = model.transcribe(ruta_audio, beam_size=5, best_of=1, **kwargs)
    except TypeError:
        result = model.transcribe(ruta_audio, **kwargs)

    base = os.path.splitext(os.path.basename(ruta_audio))[0]
    out_txt = os.path.join(TRANSCRIPCION_DIR, base + ".txt")
    out_srt = os.path.join(TRANSCRIPCION_DIR, base + ".srt")
    out_json = os.path.join(TRANSCRIPCION_DIR, base + ".json")

    raw_text = _safe_antibucle(result.get("text", ""))
    with open(out_txt, "w", encoding="utf-8") as f:
        f.write(raw_text.strip())

    segments = []
    for seg in result.get("segments", []):
        segments.append({
            "id": seg.get("id"),
            "start": seg.get("start"),
            "end": seg.get("end"),
            "text": seg.get("text")
        })
    _write_srt(segments, out_srt)

    with open(out_json, "w", encoding="utf-8") as f:
        json.dump(result, f, ensure_ascii=False, indent=2)

    print("✅ Transcripción guardada:")
    print(" - TXT :", out_txt)
    print(" - SRT :", out_srt)
    print(" - JSON:", out_json)
    print()
    return out_txt

def _pedir_audio_y_resolver(prompt_txt: str) -> str | None:
    if not ffmpeg_disponible():
        print("\n[ERROR] No se encontró FFmpeg en el PATH.")
        print("Instálalo (ej. winget install Gyan.FFmpeg) y abre nueva terminal.")
        return None

    audio_raw = input(prompt_txt)
    limpio = limpiar_ruta(audio_raw)
    print(f"[DEBUG] recibido: {repr(audio_raw)}")
    print(f"[DEBUG] limpio   : {repr(limpio)}")
    ruta = resolver_ruta_audio(limpio)
    if not ruta:
        print("[Error] No se encontró el archivo de audio.\n"
              "Sugerencias:\n"
              " - Asegúrate de soltar SOLO un archivo (no carpeta).\n"
              " - Si el nombre tiene caracteres raros, prueba renombrarlo (ej: audio.m4a).\n")
        return None
    return ruta

def transcribir_audio_fiel_auto():
    ruta = _pedir_audio_y_resolver("\nArrastra aquí tu audio y Enter (modo FIEL 1:1): ")
    if not ruta:
        return
    model_name, device, fp16, recomendado, vram_gb = elegir_config_whisper()
    cfg = _leer_config()
    if cfg.get("manual_whisper_model"):
        print(f"[INFO] Modelo manual fijado: {cfg.get('manual_whisper_model')} (sugerido para tu HW: {recomendado})")
    _transcribir_fiel(ruta, model_name, device, fp16)

def transcribir_audio_fiel_con_modelo(model_name: str):
    ruta = _pedir_audio_y_resolver("\nArrastra aquí tu audio y Enter (transcripción con modelo elegido): ")
    if not ruta:
        return
    _, device, fp16, _, _ = elegir_config_whisper()
    _transcribir_fiel(ruta, model_name, device, fp16)

# ====== PARSER SRT ======
def _parse_srt(path: str):
    """
    Devuelve lista de segmentos: [{start:'HH:MM:SS,mmm', end:'HH:MM:SS,mmm', text:'...'}]
    """
    segments = []
    if not os.path.isfile(path):
        return segments
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        lines = [ln.rstrip("\n") for ln in f]
    i = 0
    while i < len(lines):
        # índice (opcional)
        if re.match(r"^\d+\s*$", lines[i] or ""):
            i += 1
        if i >= len(lines):
            break
        # tiempo
        m = re.match(r"^(\d{2}:\d{2}:\d{2},\d{3})\s*-->\s*(\d{2}:\d{2}:\d{2},\d{3})", lines[i] or "")
        if not m:
            i += 1
            continue
        start, end = m.group(1), m.group(2)
        i += 1
        # texto (hasta línea en blanco)
        buff = []
        while i < len(lines) and lines[i].strip() != "":
            buff.append(lines[i])
            i += 1
        # saltar línea en blanco
        while i < len(lines) and lines[i].strip() == "":
            i += 1
        segments.append({"start": start, "end": end, "text": " ".join(buff).strip()})
    return segments

# ====== ACTA LITERAL (SIN CAMBIOS) ======
def _ultima_base(nombre_base: str, extension: str) -> str | None:
    base_file = os.path.join(TRANSCRIPCION_DIR, nombre_base + extension)
    if os.path.isfile(base_file):
        return base_file
    return None

def _ultima_transcripcion_txt() -> str | None:
    archivos = [f for f in os.listdir(TRANSCRIPCION_DIR) if f.lower().endswith(".txt")]
    if not archivos:
        return None
    return max([os.path.join(TRANSCRIPCION_DIR, f) for f in archivos], key=os.path.getmtime)

def generar_acta_literal():
    ultimo_txt = _ultima_transcripcion_txt()
    if not ultimo_txt:
        print("\n[Info] No se encontraron transcripciones (.txt) en la carpeta.")
        return
    base = os.path.splitext(os.path.basename(ultimo_txt))[0]
    posible_srt = _ultima_base(base, ".srt")

    # Documento
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    fecha_actual = datetime.datetime.now().strftime("%d/%m/%Y")
    hora_actual = datetime.datetime.now().strftime("%H:%M")

    def titulo(txt):
        run = doc.add_paragraph().add_run(txt)
        run.bold = True

    # Encabezado mínimo (no toca contenido hablado)
    titulo("ACTA DE REUNIÓN (LITERAL)")
    doc.add_paragraph(f"FECHA: {fecha_actual}")
    doc.add_paragraph(f"HORA DE INICIO: {hora_actual}")
    doc.add_paragraph("MODALIDAD: Presencial/Virtual")
    doc.add_paragraph("---")

    # Bloque de transcripción ÍNTEGRA
    titulo("TRANSCRIPCIÓN ÍNTEGRA (sin modificaciones)")
    if posible_srt:
        segs = _parse_srt(posible_srt)
        if segs:
            for s in segs:
                # No modificar texto ni puntuación
                p = doc.add_paragraph()
                p.add_run(f"[{s['start']} - {s['end']}] ").bold = True
                p.add_run(s["text"])
        else:
            # fallback a TXT si algo pasa con el SRT
            with open(ultimo_txt, "r", encoding="utf-8") as f:
                for ln in f:
                    doc.add_paragraph(ln.rstrip("\n"))
    else:
        # No hay SRT -> usar TXT tal cual, en líneas
        with open(ultimo_txt, "r", encoding="utf-8") as f:
            for ln in f:
                doc.add_paragraph(ln.rstrip("\n"))

    # Cierre (no altera contenido hablado)
    doc.add_paragraph("---")
    doc.add_paragraph(f"HORA DE CIERRE: {datetime.datetime.now().strftime('%H:%M')}")

    nombre_archivo = base + "_acta_literal.docx"
    ruta_salida = os.path.join(ACTAS_DIR, nombre_archivo)
    doc.save(ruta_salida)
    print(f"\n✅ Acta LITERAL generada en Word: {ruta_salida}\n")

# ====== IA LOCAL: SOLO ENCABEZADO + TRANSCRIPCIÓN ÍNTEGRA ======
def _normalizar_base_url(u: str) -> str:
    u = u.strip().rstrip("/")
    return u if u.endswith("/v1") else u + "/v1"

def _listar_modelos_lm(base_url: str, api_key: str) -> list[str]:
    try:
        url = base_url + "/models"
        r = requests.get(url, headers={"Authorization": f"Bearer {api_key}"}, timeout=10)
        r.raise_for_status()
        data = r.json()
        modelos = [it.get("id") for it in data.get("data", []) if it.get("id")]
        return modelos
    except Exception as e:
        print("[AVISO] No se pudieron listar modelos:", e)
        return []

def _elegir_modelo_lm_interactivo(modelos: list[str]) -> str:
    if not modelos:
        m = input("Modelo LM (texto, ej. openai/gpt-oss-20b): ").strip()
        return m or "openai/gpt-oss-20b"
    print("\nModelos disponibles en el servidor:")
    for i, m in enumerate(modelos, start=1):
        print(f"  {i}. {m}")
    sel = input("Elige modelo por número (o Enter para 'openai/gpt-oss-20b'): ").strip()
    if not sel:
        return "openai/gpt-oss-20b"
    try:
        idx = int(sel)
        if 1 <= idx <= len(modelos):
            return modelos[idx-1]
    except Exception:
        pass
    print("[AVISO] Entrada no válida. Usando 'openai/gpt-oss-20b'.")
    return "openai/gpt-oss-20b"

def generar_acta_con_ia_local_literal():
    ultimo_txt = _ultima_transcripcion_txt()
    if not ultimo_txt:
        print("\n[Info] No se encontraron transcripciones (.txt) en la carpeta.")
        return
    with open(ultimo_txt, "r", encoding="utf-8") as f:
        transcrip = f.read()  # ¡sin tocar!

    base = os.path.splitext(os.path.basename(ultimo_txt))[0]
    posible_srt = _ultima_base(base, ".srt")

    cfg = _leer_config()
    base_url = cfg.get("lm_base_url") or ""
    print("\n=== ACTA (IA local SOLO ENCABEZADO) + TRANSCRIPCIÓN ÍNTEGRA ===")
    print("Ejemplos: http://localhost:1234  o  http://<IP-PC>:1234")
    inp = input(f"URL del servidor [{base_url or 'http://localhost:1234'}]: ").strip()
    if not inp:
        inp = base_url or "http://localhost:1234"
    base_url_norm = _normalizar_base_url(inp)
    cfg["lm_base_url"] = inp
    _guardar_config(cfg)

    api_key = os.environ.get("LMSTUDIO_API_KEY", "lm-studio")
    modelos = _listar_modelos_lm(base_url_norm, api_key)
    modelo_lm = _elegir_modelo_lm_interactivo(modelos)

    # Pedimos SOLO encabezado/metadata. La transcripción la insertamos nosotros tal cual.
    mensajes = [
        {
            "role": "system",
            "content": (
                "Eres un asistente que genera ÚNICAMENTE el ENCABEZADO de un acta "
                "(fecha, hora inicio, lugar, modalidad, asistentes si se infieren, orden del día si existe). "
                "NUNCA reescribas, resumas ni modifiques la transcripción. "
                "No incluyas la transcripción en tu respuesta. "
                "Formato breve en español, estilo profesional."
            )
        },
        {
            "role": "user",
            "content": (
                "A partir de la siguiente transcripción, devuelve SOLO el encabezado profesional del acta. "
                "NO incluyas la transcripción, NO cambies ni resumas su contenido. "
                "Transcripción:\n\n"
                f"{transcrip}"
            )
        }
    ]

    payload = {
        "model": modelo_lm,
        "messages": mensajes,
        "temperature": 0.1,
        "max_tokens": 800
    }

    encabezado = None
    try:
        url = base_url_norm + "/chat/completions"
        r = requests.post(
            url,
            headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
            json=payload,
            timeout=(10, 180)
        )
        r.raise_for_status()
        data = r.json()
        encabezado = data.get("choices", [{}])[0].get("message", {}).get("content", "")
        if not encabezado:
            print("[AVISO] La IA devolvió contenido vacío. Se creará un encabezado mínimo.")
    except requests.exceptions.ReadTimeout:
        print("[AVISO] ReadTimeout al consultar la IA. Se creará un encabezado mínimo.")
    except requests.exceptions.ConnectionError:
        print("[AVISO] No se pudo conectar con la IA local. Se creará un encabezado mínimo.")
    except Exception as e:
        print("[AVISO] Error consultando IA local:", repr(e))

    # Crear DOCX
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    def titulo(txt):
        run = doc.add_paragraph().add_run(txt)
        run.bold = True

    titulo("ACTA DE REUNIÓN (ENCABEZADO por IA local) + TRANSCRIPCIÓN ÍNTEGRA")

    if encabezado:
        for bloque in encabezado.split("\n"):
            doc.add_paragraph(bloque)
    else:
        fecha_actual = datetime.datetime.now().strftime("%d/%m/%Y")
        hora_actual = datetime.datetime.now().strftime("%H:%M")
        doc.add_paragraph(f"FECHA: {fecha_actual}")
        doc.add_paragraph(f"HORA DE INICIO: {hora_actual}")
        doc.add_paragraph("MODALIDAD: Presencial/Virtual")

    doc.add_paragraph("---")
    titulo("TRANSCRIPCIÓN ÍNTEGRA (sin modificaciones)")

    if posible_srt and os.path.isfile(posible_srt):
        segs = _parse_srt(posible_srt)
        if segs:
            for s in segs:
                p = doc.add_paragraph()
                p.add_run(f"[{s['start']} - {s['end']}] ").bold = True
                p.add_run(s["text"])
        else:
            for ln in transcrip.splitlines():
                doc.add_paragraph(ln)
    else:
        for ln in transcrip.splitlines():
            doc.add_paragraph(ln)

    doc.add_paragraph("---")
    doc.add_paragraph(f"HORA DE CIERRE: {datetime.datetime.now().strftime('%H:%M')}")

    nombre_archivo = base + "_acta_lm_encabezado_y_transcripcion.docx"
    ruta_salida = os.path.join(ACTAS_DIR, nombre_archivo)
    doc.save(ruta_salida)
    print(f"\n✅ Acta generada: {ruta_salida}\n")

# ====== MENÚ MANUAL DE MODELOS ======
def menu_elegir_modelo_y_transcribir():
    print("\n=== Modelos Whisper disponibles ===")
    for i, (name, vram, nota) in enumerate(WHISPER_MODELOS, start=1):
        print(f"{i}. {name:9s}  | {vram:25s} | {nota}")

    sel = input("\nElige un modelo por número: ").strip()
    try:
        idx = int(sel)
        if not (1 <= idx <= len(WHISPER_MODELOS)):
            raise ValueError
    except Exception:
        print("[Error] Opción inválida.")
        return

    elegido = WHISPER_MODELOS[idx-1][0]
    print(f"\n[INFO] Elegido: {elegido}. Se iniciará la transcripción inmediatamente.")
    transcribir_audio_fiel_con_modelo(elegido)

    # Preguntar si mantener ese modelo como predeterminado
    _, _, _, sugerido, _ = elegir_config_whisper()
    print(f"\nModelo sugerido para tu hardware: {sugerido}")
    resp = input(f"¿Mantener '{elegido}' como predeterminado (S/N)? Si NO, volverá al sugerido [{sugerido}]: ").strip().lower()
    cfg = _leer_config()
    if resp == "s":
        cfg["manual_whisper_model"] = elegido
        _guardar_config(cfg)
        print(f"[OK] Modelo manual fijado: {elegido}")
    else:
        if "manual_whisper_model" in cfg:
            cfg.pop("manual_whisper_model", None)
            _guardar_config(cfg)
        print("[OK] Se usará el modelo sugerido automáticamente.")

# ====== MAIN ======
def main():
    while True:
        print("\n=== GENERADOR DE ACTAS AUTOMÁTICAS ===")
        print("1. Transcribir audio con Whisper (FIEL 1:1, auto/manual vigente)")
        print("2. Generar ACTA LITERAL (Word, sin cambios)")
        print("3. Ver configuración detectada")
        print("4. Elegir modelo Whisper (por número) y transcribir ahora")
        print("5. ACTA con IA local: encabezado IA + transcripción ÍNTEGRA")
        print("6. Salir")
        opcion = input("\nSelecciona una opción (1/2/3/4/5/6): ").strip()

        if opcion == "1":
            transcribir_audio_fiel_auto()
        elif opcion == "2":
            generar_acta_literal()
        elif opcion == "3":
            mostrar_config()
        elif opcion == "4":
            menu_elegir_modelo_y_transcribir()
        elif opcion == "5":
            generar_acta_con_ia_local_literal()
        elif opcion == "6":
            print("\nHasta pronto.")
            break
        else:
            print("\n[Error] Opción no válida. Intenta de nuevo.")

if __name__ == "__main__":
    main()
